import os
import random
import warnings
import docx
from fpdf import FPDF
from faker import Faker

#suppress warnings
warnings.filterwarnings("ignore")

#initialize faker with indian locale for realistic names and cities
fake = Faker('en_IN')

#define skill pools to create a mix of good, average, and terrible matches for the jd
HIGH_MATCH_SKILLS = ["Python", "Deep Learning", "PyTorch", "TensorFlow", "NLP", "Computer Vision", "FastAPI", "Docker", "Kubernetes", "AWS"]
MEDIUM_MATCH_SKILLS = ["Python", "SQL", "Machine Learning", "Scikit-Learn", "Pandas", "Numpy", "Data Analysis", "Tableau"]
LOW_MATCH_SKILLS = ["JavaScript", "React.js", "Node.js", "HTML", "CSS", "UI/UX", "Java", "Spring Boot", "HR", "Sales", "Marketing"]

def generate_candidate_profile() -> dict:
    """
    generates a random candidate profile dictating their name, experience, 
    skills, and match quality to ensure a diverse test dataset.
    
    returns:
        dict: a dictionary containing the candidate's generated data.
    """
    profile_type = random.choice(["high", "medium", "low"])
    
    #determine experience and skills based on profile type
    if profile_type == "high":
        #3 to 7 years of experience, mostly highly relevant skills
        exp_years = random.randint(3, 7)
        skills = random.sample(HIGH_MATCH_SKILLS, random.randint(6, 9))
        title = "AI/ML Engineer"
    elif profile_type == "medium":
        #1 to 4 years of experience, partial skill overlap
        exp_years = random.randint(1, 4)
        skills = random.sample(MEDIUM_MATCH_SKILLS, random.randint(5, 8))
        title = "Data Analyst / Scientist"
    else:
        #random experience, completely irrelevant skills
        exp_years = random.randint(2, 8)
        skills = random.sample(LOW_MATCH_SKILLS, random.randint(4, 7))
        title = random.choice(["Frontend Developer", "Backend Engineer", "Marketing Executive"])
        
    #calculate realistic start year based on experience
    start_year = 2024 - exp_years
    
    return {
        "name": fake.name(),
        "email": fake.email(),
        "location": fake.city(),
        "title": title,
        "exp_years": exp_years,
        "start_year": start_year,
        "skills": ", ".join(skills),
        "company": fake.company()
    }

def create_docx_resume(profile: dict, file_path: str):
    """
    creates a .docx resume using python-docx.
    
    args:
        profile (dict): the candidate data.
        file_path (str): the output path for the .docx file.
    """
    doc = docx.Document()
    
    #add header
    doc.add_heading(profile["name"], 0)
    doc.add_paragraph(f"Location: {profile['location']} | Email: {profile['email']}")
    
    #add summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(f"Experienced {profile['title']} with a strong background in technology and delivering high-quality solutions.")
    
    #add experience
    doc.add_heading("Experience", level=1)
    doc.add_heading(f"{profile['title']} | {profile['company']}", level=2)
    doc.add_paragraph(f"January {profile['start_year']} - Present ({profile['exp_years']} years of experience)")
    doc.add_paragraph("- Developed and maintained core business applications.\n- Collaborated with cross-functional teams to deliver projects on time.\n- Optimized system performance and scalability.")
    
    #add skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(profile["skills"])
    
    #add education
    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"Bachelor of Technology\nGraduated: {profile['start_year'] - 1}")
    
    doc.save(file_path)

def create_pdf_resume(profile: dict, file_path: str):
    """
    creates a .pdf resume using fpdf.
    
    args:
        profile (dict): the candidate data.
        file_path (str): the output path for the .pdf file.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    #add header
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(200, 10, txt=profile["name"], ln=True, align="L")
    
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 6, txt=f"Location: {profile['location']} | Email: {profile['email']}", ln=True)
    pdf.ln(5)
    
    #add summary
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 8, txt="Professional Summary", ln=True)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 6, txt=f"Experienced {profile['title']} with a strong background in technology and delivering high-quality solutions.")
    pdf.ln(5)
    
    #add experience
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 8, txt="Experience", ln=True)
    pdf.set_font("Arial", style="B", size=11)
    pdf.cell(200, 6, txt=f"{profile['title']} | {profile['company']}", ln=True)
    pdf.set_font("Arial", style="I", size=11)
    pdf.cell(200, 6, txt=f"January {profile['start_year']} - Present ({profile['exp_years']} years of experience)", ln=True)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 6, txt="- Developed and maintained core business applications.\n- Collaborated with cross-functional teams to deliver projects on time.\n- Optimized system performance and scalability.")
    pdf.ln(5)
    
    #add skills
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 8, txt="Skills", ln=True)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 6, txt=profile["skills"])
    pdf.ln(5)
    
    #add education
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 8, txt="Education", ln=True)
    pdf.set_font("Arial", size=11)
    pdf.cell(200, 6, txt="Bachelor of Technology", ln=True)
    pdf.cell(200, 6, txt=f"Graduated: {profile['start_year'] - 1}", ln=True)
    
    #save the pdf
    pdf.output(file_path)

def generate_dataset(num_resumes: int = 30):
    """
    main function to generate a batch of dummy resumes, split evenly 
    between .docx and .pdf formats.
    
    args:
        num_resumes (int): total number of resumes to generate.
    """
    #ensure the output directory exists
    output_dir = os.path.join("data", "raw_resumes")
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"generating {num_resumes} dummy resumes in '{output_dir}'...")
    
    for i in range(num_resumes):
        profile = generate_candidate_profile()
        
        #clean the name to use as a filename
        safe_name = profile["name"].replace(" ", "_").lower()
        
        #alternate between pdf and docx
        if i % 2 == 0:
            file_name = f"resume_{safe_name}.pdf"
            file_path = os.path.join(output_dir, file_name)
            create_pdf_resume(profile, file_path)
        else:
            file_name = f"resume_{safe_name}.docx"
            file_path = os.path.join(output_dir, file_name)
            create_docx_resume(profile, file_path)
            
    print("generation complete. dataset is ready for testing.")

if __name__ == "__main__":
    #generate exactly 30 resumes as requested
    generate_dataset(30)